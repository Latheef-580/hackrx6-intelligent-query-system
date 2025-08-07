import os
import tempfile
import logging
from typing import Dict, List, Any
import aiohttp
import asyncio
from urllib.parse import urlparse
import mimetypes

# Document processing libraries
import PyPDF2
from docx import Document
import email
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import re

from app.config import settings

logger = logging.getLogger(__name__)

class DocumentChunk:
    def __init__(self, text: str, metadata: Dict[str, Any]):
        self.text = text
        self.metadata = metadata
        self.chunk_id = metadata.get('chunk_id', '')

class ProcessedDocument:
    def __init__(self, doc_url: str, chunks: List[DocumentChunk], metadata: Dict[str, Any]):
        self.doc_url = doc_url
        self.chunks = chunks
        self.metadata = metadata

class DocumentProcessor:
    def __init__(self):
        self.chunk_size = settings.CHUNK_SIZE
        self.chunk_overlap = settings.CHUNK_OVERLAP
        self.session = None
    
    async def _get_session(self):
        """Get or create aiohttp session"""
        if self.session is None or self.session.closed:
            timeout = aiohttp.ClientTimeout(total=30)
            self.session = aiohttp.ClientSession(timeout=timeout)
        return self.session
    
    async def _download_document(self, doc_url: str) -> tuple[bytes, str]:
        """Download document from URL and determine file type"""
        try:
            session = await self._get_session()
            async with session.get(doc_url) as response:
                if response.status != 200:
                    raise Exception(f"Failed to download document: HTTP {response.status}")
                
                content = await response.read()
                content_type = response.headers.get('content-type', '').lower()
                
                # Determine file extension from URL or content type
                parsed_url = urlparse(doc_url)
                file_extension = os.path.splitext(parsed_url.path)[1].lower()
                
                if not file_extension:
                    # Guess from content type
                    if 'pdf' in content_type:
                        file_extension = '.pdf'
                    elif 'word' in content_type or 'docx' in content_type:
                        file_extension = '.docx'
                    elif 'email' in content_type or 'message' in content_type:
                        file_extension = '.eml'
                    else:
                        # Try to detect from content
                        if content.startswith(b'%PDF'):
                            file_extension = '.pdf'
                        elif b'PK' in content[:4]:  # ZIP-based format (DOCX)
                            file_extension = '.docx'
                        else:
                            file_extension = '.txt'
                
                logger.info(f"Downloaded document: {len(content)} bytes, type: {file_extension}")
                return content, file_extension
                
        except Exception as e:
            logger.error(f"Error downloading document from {doc_url}: {str(e)}")
            raise
    
    def _extract_text_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            # Use BytesIO to avoid file system issues
            import io
            pdf_stream = io.BytesIO(content)
            
            text = ""
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
            
            return text.strip()
                
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            # Use BytesIO to avoid file system issues
            import io
            docx_stream = io.BytesIO(content)
            
            doc = Document(docx_stream)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += f"TABLE: {row_text}\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _extract_text_from_email(self, content: bytes) -> str:
        """Extract text from email content"""
        try:
            # Try to decode as string first
            try:
                email_str = content.decode('utf-8')
            except UnicodeDecodeError:
                email_str = content.decode('latin-1')
            
            msg = email.message_from_string(email_str)
            text = ""
            
            # Extract headers
            text += f"From: {msg.get('From', 'Unknown')}\n"
            text += f"To: {msg.get('To', 'Unknown')}\n"
            text += f"Subject: {msg.get('Subject', 'No Subject')}\n"
            text += f"Date: {msg.get('Date', 'Unknown')}\n\n"
            
            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            try:
                                text += payload.decode('utf-8') + "\n"
                            except UnicodeDecodeError:
                                text += payload.decode('latin-1', errors='ignore') + "\n"
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    try:
                        text += payload.decode('utf-8')
                    except UnicodeDecodeError:
                        text += payload.decode('latin-1', errors='ignore')
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting text from email: {str(e)}")
            # Fallback: treat as plain text
            try:
                return content.decode('utf-8', errors='ignore')
            except Exception:
                return content.decode('latin-1', errors='ignore')
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove page numbers and headers/footers patterns
        text = re.sub(r'\n--- Page \d+ ---\n', '\n', text)
        text = re.sub(r'\nPage \d+ of \d+\n', '\n', text)
        
        # Clean up special characters
        text = text.replace('\x00', '')
        text = text.replace('\ufffd', '')
        
        return text.strip()
    
    def _chunk_text(self, text: str, doc_url: str) -> List[DocumentChunk]:
        """Split text into overlapping chunks"""
        chunks = []
        words = text.split()
        
        if len(words) <= self.chunk_size:
            # Document is small enough to be one chunk
            chunk = DocumentChunk(
                text=text,
                metadata={
                    'chunk_id': f"{doc_url}_chunk_0",
                    'doc_url': doc_url,
                    'chunk_index': 0,
                    'total_chunks': 1,
                    'word_count': len(words)
                }
            )
            chunks.append(chunk)
        else:
            # Split into overlapping chunks
            chunk_index = 0
            start = 0
            
            while start < len(words):
                end = min(start + self.chunk_size, len(words))
                chunk_words = words[start:end]
                chunk_text = ' '.join(chunk_words)
                
                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata={
                        'chunk_id': f"{doc_url}_chunk_{chunk_index}",
                        'doc_url': doc_url,
                        'chunk_index': chunk_index,
                        'word_count': len(chunk_words),
                        'start_word': start,
                        'end_word': end
                    }
                )
                chunks.append(chunk)
                
                # Move start position with overlap
                start += self.chunk_size - self.chunk_overlap
                chunk_index += 1
        
        # Update total chunks count
        for chunk in chunks:
            chunk.metadata['total_chunks'] = len(chunks)
        
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
    
    async def process_document(self, doc_url: str) -> ProcessedDocument:
        """Main method to process a document from URL"""
        try:
            logger.info(f"Processing document: {doc_url}")
            
            # Download document
            content, file_extension = await self._download_document(doc_url)
            
            # Extract text based on file type
            if file_extension == '.pdf':
                raw_text = self._extract_text_from_pdf(content)
            elif file_extension == '.docx':
                raw_text = self._extract_text_from_docx(content)
            elif file_extension in ['.eml', '.email']:
                raw_text = self._extract_text_from_email(content)
            else:
                # Fallback: treat as plain text
                try:
                    raw_text = content.decode('utf-8')
                except UnicodeDecodeError:
                    raw_text = content.decode('latin-1', errors='ignore')
            
            if not raw_text or len(raw_text.strip()) < 10:
                raise Exception("No meaningful text could be extracted from the document")
            
            # Clean text
            cleaned_text = self._clean_text(raw_text)
            
            # Create chunks
            chunks = self._chunk_text(cleaned_text, doc_url)
            
            # Create document metadata
            metadata = {
                'doc_url': doc_url,
                'file_type': file_extension,
                'total_chars': len(cleaned_text),
                'total_words': len(cleaned_text.split()),
                'total_chunks': len(chunks),
                'processed_at': asyncio.get_event_loop().time()
            }
            
            processed_doc = ProcessedDocument(doc_url, chunks, metadata)
            
            logger.info(f"Successfully processed document: {len(chunks)} chunks, {metadata['total_words']} words")
            return processed_doc
            
        except Exception as e:
            logger.error(f"Error processing document {doc_url}: {str(e)}")
            raise
    
    async def close(self):
        """Close aiohttp session"""
        if self.session and not self.session.closed:
            await self.session.close()